from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


class ReportWriterTool:

    def __init__(
        self,
        output_directory: str = "outputs/reports"
    ):
        self.output_directory = Path(
            output_directory
        )

        self.output_directory.mkdir(
            parents=True,
            exist_ok=True
        )

    def write_report(
        self,
        title: str,
        content: str,
        filename: Optional[str] = None
    ) -> dict:

        # =====================================================
        # VALIDATION
        # =====================================================

        if not title or not title.strip():
            raise ValueError(
                "Report title cannot be empty."
            )

        if not content or not content.strip():
            raise ValueError(
                "Report content cannot be empty."
            )

        title = title.strip()
        content = content.strip()

        # =====================================================
        # FILE NAME
        # =====================================================

        if not filename:

            filename = self._create_filename(
                title
            )

        filename = self._sanitize_filename(
            filename
        )

        if not filename.lower().endswith(".docx"):
            filename += ".docx"

        output_path = (
            self.output_directory / filename
        )

        # =====================================================
        # CREATE DOCUMENT
        # =====================================================

        document = Document()

        # -----------------------------------------------------
        # Default font
        # -----------------------------------------------------

        styles = document.styles

        styles["Normal"].font.name = "Arial"
        styles["Normal"].font.size = Pt(11)

        # -----------------------------------------------------
        # Title
        # -----------------------------------------------------

        title_paragraph = document.add_paragraph()

        title_paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        title_run = title_paragraph.add_run(
            title
        )

        title_run.bold = True
        title_run.font.size = Pt(18)

        # -----------------------------------------------------
        # Report content
        # -----------------------------------------------------

        self._add_content(
            document,
            content
        )

        # =====================================================
        # SAVE
        # =====================================================

        document.save(
            str(output_path)
        )

        # =====================================================
        # RETURN STRUCTURED RESULT
        # =====================================================

        return {
            "status": "success",
            "file_type": "docx",
            "filename": filename,
            "path": str(
                output_path.resolve()
            ),
            "title": title,
            "content_length": len(content),
        }

    # =========================================================
    # CONTENT PROCESSING
    # =========================================================

    def _add_content(
        self,
        document: Document,
        content: str
    ):

        lines = content.splitlines()

        for line in lines:

            line = line.strip()

            if not line:
                continue

            # -------------------------------------------------
            # Markdown headings
            # -------------------------------------------------

            if line.startswith("### "):

                paragraph = document.add_heading(
                    line[4:].strip(),
                    level=3
                )

                continue

            if line.startswith("## "):

                paragraph = document.add_heading(
                    line[3:].strip(),
                    level=2
                )

                continue

            if line.startswith("# "):

                paragraph = document.add_heading(
                    line[2:].strip(),
                    level=1
                )

                continue

            # -------------------------------------------------
            # Bullet points
            # -------------------------------------------------

            if line.startswith("- "):

                document.add_paragraph(
                    line[2:].strip(),
                    style="List Bullet"
                )

                continue

            if line.startswith("* "):

                document.add_paragraph(
                    line[2:].strip(),
                    style="List Bullet"
                )

                continue

            # -------------------------------------------------
            # Numbered list
            # -------------------------------------------------

            if self._is_numbered_item(line):

                text = self._remove_number(
                    line
                )

                document.add_paragraph(
                    text,
                    style="List Number"
                )

                continue

            # -------------------------------------------------
            # Normal paragraph
            # -------------------------------------------------

            document.add_paragraph(
                line
            )

    # =========================================================
    # HELPERS
    # =========================================================

    def _is_numbered_item(
        self,
        text: str
    ) -> bool:

        parts = text.split(
            ".",
            1
        )

        if len(parts) != 2:
            return False

        return parts[0].strip().isdigit()

    def _remove_number(
        self,
        text: str
    ) -> str:

        return text.split(
            ".",
            1
        )[1].strip()

    def _sanitize_filename(
        self,
        filename: str
    ) -> str:

        invalid_chars = '<>:"/\\|?*'

        for char in invalid_chars:
            filename = filename.replace(
                char,
                "_"
            )

        return filename.strip()

    def _create_filename(
        self,
        title: str
    ) -> str:

        filename = self._sanitize_filename(
            title
        )

        return f"{filename}.docx"